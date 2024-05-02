import matplotlib.pyplot as plt
import matplotlib
import scipy.fftpack
from scipy.signal import find_peaks
import numpy as np
import csv
import pandas as pd
from pandas import DataFrame
import pandas_profiling
from pandas_profiling import ProfileReport
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.graph_objects as go
from plotly.offline import plot
import plotly.io as pio
from datetime import datetime
import datetime as dt
from datetime import timedelta
import math
import operator
from operator import methodcaller
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
import pathlib
import sys

def main(csv_file, dir_path):
    filename = os.path.splitext(os.path.basename(csv_file))[0]  # Extract filename without extension
    Open_Directory = dir_path + os.path.sep  # Extract the directory of the CSV file

    df = pd.read_csv(csv_file)
    if os.path.isfile(csv_file):
        print("File is present")  # check file address is correct
    else:
        print("file cannot be found")

    # ## file data

    now = datetime.now()
    DataTime = now.strftime("%Y-%m-%d, %H-%M")

    DisplayGraphsInBrowser = True
    DisplayGraphsInSeries = False
    SaveData = True

    if SaveData == True:
        newpath = Open_Directory + str(DataTime) + "_" + filename
        if not os.path.exists(newpath):
            os.makedirs(newpath)

        print(newpath)

    # ## import raw data from file

    # TestData = df.drop(list(range(0)))   # drops unwanted rows until dataseries names
    TestData = df
    TestData

    # # Global DataFrame

    # Creation of single GLobal DataFrame with an intial sort
    GlobalDF = DataFrame(TestData)  # , columns=[""].astype(float))
    # GlobalDF = GlobalDF[sorted(GlobalDF)]
    ##########################################################

    # Sets up Test Time
    GlobalDF["Timestamp"] = pd.to_datetime(GlobalDF["Timestamp"])  # , format='%d/%m/%Y %H:%M')
    # Convert the 'Timestamp' column to datetime
    GlobalDF['Timestamp'] = pd.to_datetime(GlobalDF['Timestamp'])
    # Convert the 'Timestamp' column to seconds
    # GlobalDF['TestTime_(ms)'] = GlobalDF['Timestamp'].apply(lambda x: x.timestamp())
    GlobalDF['TestTime_(ms)'] = (GlobalDF['Timestamp'] - GlobalDF['Timestamp'].iloc[0]).dt.total_seconds()

    GlobalDF

    ## Data Analysis

    ## Using the inital imported data further analysis can be performed by simply creating new GlobalDF

    ## find resolution of timeseries
    ResolutionS = (GlobalDF['TestTime_(ms)'].iloc[2] - GlobalDF['TestTime_(ms)'].iloc[1])
    print(ResolutionS)

    try:
        # Map LED state to 1 or 0
        GlobalDF["LED_Log"] = GlobalDF["LED state"].apply(lambda x: 1.0 if x == 'on' else 0.0)
    except KeyError:
        print("Dataseries not included.")

    try:
        # Convert Thumbstick_Xvalue to position
        GlobalDF["TVC_X_Position"] = (GlobalDF["Thumbstick_Xvalue"] - 1822) / 182.2
    except KeyError:
        print("Dataseries not included.")

    try:
        # Convert Thumbstick_Xvalue to position
        GlobalDF["TVC_Y_Position"] = (GlobalDF["Thumbstick_Yvalue"] - 1797) / 1797
    except KeyError:
        print("Dataseries not included.")

    GlobalDF

    # # Plot_2D Interactive

    # Setup Plot
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    fig.update_layout(legend_orientation="v", title="<b>" + str(filename) + "<b>", title_font_size=28,
                      legend_font_size=22)  # Adjust the legend font size here
    # fig.update_layout(height=1920, width = 1080, legend_orientation="v", title="<b>"+str(file)+"<b>", title_font_size = 24)
    fig.update_yaxes(title_text="Amplitude", title_font_size=24, row=1, col=1)
    fig.update_yaxes(title_text="", secondary_y=True, row=1, col=1)
    fig.update_xaxes(title_text="Time", title_font_size=24, row=1, col=1)
    fig.update_layout(legend_borderwidth=4)

    # Plot GlobalDF
    ## group for clarity
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["Thumbstick_Xvalue"], mode='lines',
                                 name="Thumbstick_Xvalue"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["Thumbstick_Yvalue"], mode='lines',
                                 name="Thumbstick_Yvalue"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["dataseries3"], mode='markers',
                                 name="dataseries3"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["dataseries4"], mode='lines',
                                 name="dataseries4"))
    except:
        print("dataseries not incluced")
    ##########################################################
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["TVC_X_Position"], mode='lines',
                                 name="TVC_X_Position"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["TVC_Y_Position"], mode='lines',
                                 name="TVC_Y_Position"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["TVC_Command_Xvalue"], mode='lines',
                                 name="TVC_Command_Xvalue"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["TVC_Command_Yvalue"], mode='lines',
                                 name="TVC_Command_Yvalue"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["LED_Log"], mode='markers',
                                 name="LED_Log"))
    except:
        print("dataseries not incluced")
    try:
        fig.add_trace(go.Scatter(x=GlobalDF["TestTime_(ms)"], y=GlobalDF["Temperature (Celius)"], mode='markers',
                                 name="Temperature (Celius)"))
    except:
        print("dataseries not incluced")
    ##########################################################

    # Display Graph
    if DisplayGraphsInBrowser == True:
        plot(fig, show_link=True, filename=newpath + "\\" + "2D_Test_Plot" + ".html")
        pio.write_image(fig, newpath + "\\" + '2D_Test_Plot.png', width=1920, height=1080, format='png')
    if DisplayGraphsInSeries == True:
        print("display check")
        fig.show

    # # GlobalDF Mean, Min, Max, STD

    # use this section to final the mean, min, max and standard diiviation of all dataseries
    # include a final sorting

    # =================== Prints Data Summary Table =============================#
    DataSummary = GlobalDF
    DataSummary = GlobalDF[sorted(DataSummary)]
    DataSummary = DataSummary.describe()
    # DataSummary=DataSummary.describe(percentiles=[.5])
    try:
        DataSummary = DataSummary.drop(columns=("x"))
    except:
        print("dataseries not inluced")
    DataSummary = DataSummary.round(2)
    DataSummary = DataSummary.T

    if (SaveData == True):
        DataSummary.to_csv(newpath + "\\" + "Test_Data_Summary" + ".csv")
        GlobalDF.to_csv(newpath + "\\" + "Test_Analysis_Data" + ".csv")

    DataSummary

    # ## Export Report

    # ## Create Report Document in Word

    document = Document(
        r"D:\Documents\GitHub\Sunfire-TVC-FYP_23.24\Data_Analysis_Templates\Test_Report_Template_v1.0.docx")
    # document = Document(r"C:\\Users\\jacob\\Documents\\GitHub\\Sunfire-TVC-FYP_23.24\\Data_Analysis_Templates\\Test_Report_Template_v1.0.docx")
    # the style types must be stored in the doc, word stores them in the program until they are used.
    # create a new paragraph and assign it the style, before deleting the paragraph and saving.

    document.add_heading('Test_Report' + DataTime, 0)

    p = document.add_paragraph('This is an auto-generated test report produced by: ')
    p.add_run('Jacob Lawson').bold = True

    document.add_picture(newpath + "\\" + "2D_Test_Plot.png", width=Inches(7))

    document.add_page_break()

    document.add_heading('Test Data Summary Table', level=1)
    document.add_paragraph('The following data is a summary of both the source recorded data series and the processed data',
                           style='Intense Quote')

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = document.add_table(DataSummary.shape[0] + 1, DataSummary.shape[1] + 1)
    t.style = "Table Grid"

    # add the header rows.
    for j in range(DataSummary.shape[-1]):
        t.cell(0, j + 1).text = DataSummary.columns[j]

    # include index
    t.cell(0, 0).text = "Data Series Index"
    for i in range(DataSummary.shape[0]):
        t.cell(i + 1, 0).text = DataSummary.index[i]

    # add the rest of the data frame
    for i in range(DataSummary.shape[0]):
        for j in range(DataSummary.shape[-1]):
            t.cell(i + 1, j + 1).text = str(DataSummary.values[i, j])

    # make headers and index bold
    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def make_cols_bold(*columns):
        for col in columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    make_rows_bold(t.rows[0])
    make_cols_bold(t.columns[0])

    document.add_page_break()

    document.save(newpath + "\\" + "Test_Report" + DataTime + ".docx")

    # ## Convert Word Report to PDF
    convert(newpath + "\\" + "Test_Report" + DataTime + ".docx", newpath + "\\" + "Test_Report" + DataTime + ".pdf")


if __name__ == "__main__":
    if len(sys.argv) > 2:
        csv_file = sys.argv[1]
        dir_path = sys.argv[2]
        main(csv_file, dir_path)
    else:
        print("Usage: python analysis_script.py <csv_file> <dir_path>")
